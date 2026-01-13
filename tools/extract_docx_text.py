from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(docx_path)

    lines: list[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [
                (cell.text or "").strip().replace("\n", " ") for cell in row.cells
            ]
            if any(cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip() + "\n"


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    requirements_dir = repo_root / "Requirements"

    if not requirements_dir.exists():
        raise SystemExit(f"Missing Requirements folder: {requirements_dir}")

    docx_files = sorted(requirements_dir.glob("*.docx"))
    if not docx_files:
        raise SystemExit(f"No .docx files found in: {requirements_dir}")

    for docx_path in docx_files:
        extracted = extract_docx_text(docx_path)
        out_path = requirements_dir / f"_extracted_{docx_path.stem}.txt"
        out_path.write_text(extracted, encoding="utf-8")
        print(f"Wrote {out_path.name} (chars={len(extracted)})")


if __name__ == "__main__":
    main()
